"""
试卷解析器
支持 PDF / Word / HTML 文件解析，自动拆分题目
"""
import os
import re
from dataclasses import dataclass, field
from typing import Optional

from config import SUBJECTS


@dataclass
class Question:
    """题目数据结构"""
    number: int
    q_type: str  # choice / fill / solve
    content: str = ""
    options: list = field(default_factory=list)
    answer: str = ""
    score: float = 0.0
    knowledge_points: list = field(default_factory=list)


@dataclass
class ParsedPaper:
    """解析后的试卷"""
    title: str
    subject: str
    total_score: float = 150.0
    questions: list = field(default_factory=list)


class PaperParser:
    """试卷解析器"""

    # 题目类型识别正则
    CHOICE_PATTERN = re.compile(
        r"(?:^|\n)\s*(\d+)[\.、．]\s*(.+?)(?:\n|$)", re.MULTILINE
    )
    ANSWER_PATTERN = re.compile(
        r"(?:^|\n)\s*(?:参考答案|答案|解析)\s*[:：]\s*(.+?)(?:\n|$)", re.MULTILINE
    )
    SECTION_PATTERN = re.compile(
        r"(?:一|二|三|四|五|六|七|八|九|十)[、．\.]\s*(?:选择题|填空题|解答题|计算题|实验题|选考题|必考题)"
    )

    def parse_file(self, file_path: str, subject_key: str = "math") -> ParsedPaper:
        """根据文件类型选择解析器"""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return self._parse_pdf(file_path, subject_key)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(file_path, subject_key)
        elif ext in (".html", ".htm"):
            return self._parse_html(file_path, subject_key)
        elif ext == ".txt":
            return self._parse_text(file_path, subject_key)
        else:
            return ParsedPaper(title=os.path.basename(file_path), subject=subject_key)

    def parse_text(self, text: str, subject_key: str = "math") -> ParsedPaper:
        """直接解析文本内容"""
        return self._parse_text(text, subject_key)

    def _parse_pdf(self, file_path, subject_key):
        """解析 PDF"""
        try:
            import pdfplumber
        except ImportError:
            return ParsedPaper(title=os.path.basename(file_path), subject=subject_key)

        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        title = os.path.basename(file_path)
        paper = self._parse_text(text, subject_key)
        paper.title = title
        return paper

    def _parse_docx(self, file_path, subject_key):
        """解析 Word 文件"""
        try:
            from docx import Document
        except ImportError:
            return ParsedPaper(title=os.path.basename(file_path), subject=subject_key)

        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"

        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        title = os.path.basename(file_path)
        paper = self._parse_text(text, subject_key)
        paper.title = title
        return paper

    def _parse_html(self, file_path, subject_key):
        """解析 HTML"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            return ParsedPaper(title=os.path.basename(file_path), subject=subject_key)

        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "lxml")

        # 移除 script/style
        for tag in soup(["script", "style"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        title = soup.title.string if soup.title else os.path.basename(file_path)
        paper = self._parse_text(text, subject_key)
        paper.title = title or os.path.basename(file_path)
        return paper

    def _parse_text(self, text: str, subject_key: str) -> ParsedPaper:
        """核心文本解析逻辑 - 自动拆分题目"""
        paper = ParsedPaper(
            title="",
            subject=subject_key,
            total_score=SUBJECTS.get(subject_key, {}).get("total_score", 150),
        )

        lines = text.strip().split("\n")
        current_section = "未知"
        questions = []
        current_question = None
        question_buffer = []

        # 题型分值映射（根据科目的常见配置）
        score_config = self._get_score_config(subject_key)

        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                continue

            # 识别大题标题
            section_match = self.SECTION_PATTERN.search(line_stripped)
            if section_match:
                current_section = section_match.group()
                continue

            # 识别题号开头
            q_match = re.match(r"\s*(\d{1,2})[\.、．．\)]\s*(.*)", line_stripped)
            if q_match:
                # 保存上一题
                if current_question:
                    current_question.content = "\n".join(question_buffer).strip()
                    questions.append(current_question)

                q_num = int(q_match.group(1))
                q_content = q_match.group(2)

                # 判断题型
                q_type = self._detect_question_type(
                    current_section, q_content, q_num, score_config
                )

                # 获取分值
                score = self._get_question_score(q_num, q_type, score_config)

                current_question = Question(
                    number=q_num,
                    q_type=q_type,
                    content=q_content,
                    score=score,
                )
                question_buffer = [q_content]
            elif current_question:
                # 追加到当前题目
                question_buffer.append(line_stripped)

                # 检测选项
                opt_match = re.match(r"\s*([A-D])[\.、．\)]\s*(.*)", line_stripped)
                if opt_match and current_question.q_type == "choice":
                    current_question.options.append(
                        {"label": opt_match.group(1), "text": opt_match.group(2)}
                    )

        # 最后一题
        if current_question:
            current_question.content = "\n".join(question_buffer).strip()
            questions.append(current_question)

        paper.questions = questions
        return paper

    def _detect_question_type(self, section, content, q_num, score_config):
        """根据大题标题和题号范围判断题型"""
        if "选择题" in section:
            return "choice"
        elif "填空题" in section:
            return "fill"
        elif any(kw in section for kw in ["解答", "计算", "证明", "实验", "简答"]):
            return "solve"
        elif content and any(
            kw in content for kw in ["选择", "下列", "不正确", "正确的是"]
        ):
            return "choice"
        else:
            # 根据题号范围推断
            if score_config.get("choice_end") and q_num <= score_config["choice_end"]:
                return "choice"
            elif score_config.get("fill_end") and q_num <= score_config["fill_end"]:
                return "fill"
            return "solve"

    def _get_score_config(self, subject_key):
        """获取各科目的分值配置"""
        configs = {
            "math": {
                "choice_count": 8, "choice_score": 5, "choice_end": 8,
                "fill_count": 4, "fill_score": 5, "fill_end": 12,
                "solve_count": 6, "solve_start_score": 10,
            },
            "chinese": {
                "choice_count": 10, "choice_score": 3, "choice_end": 10,
                "fill_count": 0, "fill_end": 10,
                "solve_count": 5, "solve_start_score": 20,
            },
            "english": {
                "choice_count": 15, "choice_score": 2, "choice_end": 20,
                "fill_count": 0, "fill_end": 20,
                "solve_count": 3, "solve_start_score": 25,
            },
            "physics": {
                "choice_count": 8, "choice_score": 4, "choice_end": 8,
                "fill_count": 2, "fill_score": 4, "fill_end": 10,
                "solve_count": 3, "solve_start_score": 10,
            },
            "chemistry": {
                "choice_count": 7, "choice_score": 6, "choice_end": 7,
                "fill_count": 3, "fill_score": 4, "fill_end": 10,
                "solve_count": 3, "solve_start_score": 10,
            },
            "biology": {
                "choice_count": 6, "choice_score": 6, "choice_end": 6,
                "fill_count": 4, "fill_score": 4, "fill_end": 10,
                "solve_count": 2, "solve_start_score": 15,
            },
            "history": {
                "choice_count": 16, "choice_score": 3, "choice_end": 16,
                "fill_count": 0, "fill_end": 16,
                "solve_count": 4, "solve_start_score": 13,
            },
            "geography": {
                "choice_count": 11, "choice_score": 4, "choice_end": 11,
                "fill_count": 0, "fill_end": 11,
                "solve_count": 3, "solve_start_score": 18,
            },
            "politics": {
                "choice_count": 16, "choice_score": 3, "choice_end": 16,
                "fill_count": 0, "fill_end": 16,
                "solve_count": 4, "solve_start_score": 13,
            },
        }
        return configs.get(subject_key, configs["math"])

    def _get_question_score(self, q_num, q_type, score_config):
        """根据题号和题型估算分值"""
        if q_type == "choice":
            return score_config.get("choice_score", 5)
        elif q_type == "fill":
            return score_config.get("fill_score", 5)
        else:
            return score_config.get("solve_start_score", 12)

    def extract_answers(self, text: str) -> dict:
        """从答案部分提取答案"""
        answers = {}
        lines = text.split("\n")
        in_answer_section = False

        for line in lines:
            stripped = line.strip()
            if "参考答案" in stripped or "答案" in stripped:
                in_answer_section = True
                continue

            if in_answer_section:
                match = re.match(r"(\d+)[\.、．\)]\s*(.+?)(?:\s|$)", stripped)
                if match:
                    q_num = int(match.group(1))
                    ans = match.group(2)
                    answers[q_num] = ans
                elif "解析" in stripped or "详细" in stripped:
                    break

        return answers
